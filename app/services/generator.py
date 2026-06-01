import subprocess
import tempfile
import os
import base64
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.services.parser import fmt_amount

NAVY   = RGBColor(0x1a, 0x3a, 0x5c)
BLACK  = RGBColor(0x00, 0x00, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GOLD   = "FFD700"
LIGHT_BLUE = "E8EFF8"
NOTICE_BG  = "EEF4FF"
HEADER_BG  = "1A3A5C"

LOGO_PATH = Path(__file__).parent.parent.parent / "assets" / "ybt_logo.jpg"

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, color="000000", size="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _para(cell, text: str, bold=False, color=BLACK, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    p  = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def _add_para(cell, text: str, bold=False, color=BLACK, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    p  = cell.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def generate_docx(client: Dict[str, Any], container_num: str, load_date: str, eta: str) -> bytes:
    doc = Document()

    # ── Marges ───────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # ── HEADER : logo + infos ─────────────────────────────────────────────────
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.style = "Table Grid"
    for cell in hdr_table.rows[0].cells:
        for side in ("top","left","bottom","right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))

    logo_cell = hdr_table.rows[0].cells[0]
    logo_cell.width = Inches(1.5)
    if LOGO_PATH.exists():
        logo_cell.paragraphs[0].add_run().add_picture(str(LOGO_PATH), width=Inches(1.3))
    logo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_cell = hdr_table.rows[0].cells[1]
    _para(info_cell, "YBT INTERNATIONAL OCEAN FREIGHT & LOGISTIC",
          bold=True, color=NAVY, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    for line in [
        "佛山市南海区里水镇里官路赤山段 72 号 A5 仓库",
        "No.72 Chishan Section, Li guan Road, Lishui Town, Nanhai District, Foshan City",
        "MOB: +8613145750861 / +8613612627740 / +23280983049 / +23276711790",
    ]:
        _add_para(info_cell, line, size=8, color=RGBColor(0x44,0x44,0x44))

    doc.add_paragraph()

    # ── Date + titre ──────────────────────────────────────────────────────────
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_p.add_run(f"40HQ GROUPAGE   {load_date}")
    run.bold = True; run.font.size = Pt(9)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("LOADING INVOICE")
    run.bold = True; run.font.size = Pt(14); run.font.color.rgb = NAVY

    # ── Notice ───────────────────────────────────────────────────────────────
    notice = doc.add_table(rows=1, cols=1)
    notice.style = "Table Grid"
    nc = notice.rows[0].cells[0]
    _set_cell_bg(nc, NOTICE_BG)
    _para(nc,
        "ALL INVOICE ARE EXPECTED TO BE PAID TWO WEEK BEFORE EXPECTED DELIVERY DATE",
        bold=True, color=NAVY, size=8)
    _add_para(nc,
        "KINDLY CONTACT THE NUMBERS BELOW FOR PAYMENT AND GET A RECEIPT OF ALL "
        "PAYMENT MADE IN ORDER TO CLAIM OR COLLECT YOUR GOODS", size=8)
    _add_para(nc,
        "MR IBRAHIM : +23276711790     MR EDDIE : +23280983049",
        bold=True, color=NAVY, size=8)

    doc.add_paragraph()

    # ── Tableau principal ─────────────────────────────────────────────────────
    headers = ["CUSTOMER NAME","CONTAINER #","DESCRIPTION","QTY","CBM","FREIGHT","CUSTOM","ETA"]
    widths  = [Inches(1.4), Inches(1.1), Inches(2.1), Inches(0.5), Inches(0.6),
               Inches(0.7), Inches(0.7), Inches(0.9)]

    inv = doc.add_table(rows=1, cols=8)
    inv.style = "Table Grid"

    # Header row
    for i, (cell, hdr) in enumerate(zip(inv.rows[0].cells, headers)):
        cell.width = widths[i]
        _set_cell_bg(cell, HEADER_BG)
        _set_cell_border(cell)
        _para(cell, hdr, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data row
    total_cbm   = client["total_cbm"]
    freight     = client["freight"]
    custom      = client["custom"]
    total_due   = client["total_due"]
    desc_lines  = "\n".join(f"• {it['description']}  ({it['quantity']})"
                            for it in client["items"] if float(it.get("cbm",0)) > 0)
    total_qty   = sum(int(re.sub(r"[^\d]","",str(it.get("quantity","1"))) or "1")
                      for it in client["items"])

    import re
    row = inv.add_row()
    data = [
        f"{client['name']}\n{client.get('phone','')}",
        container_num,
        desc_lines,
        str(total_qty),
        fmt_amount(total_cbm),
        f"${fmt_amount(freight)}",
        f"${fmt_amount(custom)}",
        f"Estimated\nArrival\n{eta}",
    ]
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]

    for i, (cell, text, align) in enumerate(zip(row.cells, data, aligns)):
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_border(cell)
        lines = text.split("\n")
        _para(cell, lines[0], size=8, align=align,
              color=NAVY if i in (5,6) else BLACK, bold=(i in (5,6)))
        for l in lines[1:]:
            _add_para(cell, l, size=8, align=align)

    # 2 lignes vides
    for _ in range(2):
        er = inv.add_row()
        for cell in er.cells:
            _set_cell_border(cell)

    doc.add_paragraph()

    # ── Récapitulatif ─────────────────────────────────────────────────────────
    # Tableau vide 1 col pour aligner à droite
    outer = doc.add_table(rows=1, cols=2)
    outer.style = "Table Grid"
    for cell in outer.rows[0].cells:
        for side in ("top","left","bottom","right"):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBorders.append(el)
            tcPr.append(tcBorders)

    outer.rows[0].cells[0].width = Inches(4.5)
    sum_cell = outer.rows[0].cells[1]
    sum_cell.width = Inches(3.5)

    sum_tbl = sum_cell.add_table(rows=5, cols=2)
    sum_tbl.style = "Table Grid"
    sum_data = [
        ("Total CBM",     fmt_amount(total_cbm) + " m³", False, LIGHT_BLUE),
        ("Tarif ($/CBM)", f"${fmt_amount(client.get('rate', 280))}", False, LIGHT_BLUE),
        ("Freight",       f"${fmt_amount(freight)}",  False, LIGHT_BLUE),
        ("Custom",        f"${fmt_amount(custom)}",   False, LIGHT_BLUE),
        ("TOTAL À PAYER", f"${fmt_amount(total_due)}", True,  GOLD),
    ]
    for row_data, row_obj in zip(sum_data, sum_tbl.rows):
        label, value, bold, bg = row_data
        for cell in row_obj.cells:
            _set_cell_border(cell)
            _set_cell_bg(cell, bg)
        _para(row_obj.cells[0], label, bold=True, size=8,
              color=NAVY if bold else BLACK)
        _para(row_obj.cells[1], value, bold=bold, size=8,
              color=NAVY if bold else BLACK, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("YBT International Ocean Freight & Logistic — Foshan, China")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x88,0x88,0x88)

    # ── Sauvegarde temporaire ─────────────────────────────────────────────────
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        return tmp.name

def docx_to_pdf(docx_path: str) -> str:
    """Convertit un .docx en PDF via LibreOffice."""
    out_dir = str(Path(docx_path).parent)
    subprocess.run([
        "soffice", "--headless", "--convert-to", "pdf",
        docx_path, "--outdir", out_dir
    ], check=True, capture_output=True)
    pdf_path = docx_path.replace(".docx", ".pdf")
    if not Path(pdf_path).exists():
        raise FileNotFoundError(f"PDF non généré : {pdf_path}")
    return pdf_path

def generate_invoice_pdf(client: Dict[str, Any], container_num: str,
                          load_date: str, eta: str) -> str:
    """Pipeline complet : génère le PDF et retourne son chemin."""
    docx_path = generate_docx(client, container_num, load_date, eta)
    pdf_path  = docx_to_pdf(docx_path)
    os.unlink(docx_path)  # nettoyage
    return pdf_path
